from docx import Document
import pandas as pd

class CheckFiles:
    def __init__(self, f1, f2):
        self.lista1, self.lista2, self.lista_diferenca = [], [], []
        self.f1 = f1
        self.f2 = f2
        self.document1 = Document(f1).paragraphs
        self.document2 = Document(f2).paragraphs
        pd.set_option('display.max_colwidth', None)


    def verifica_tamanho(self):
        """Checa se a quantidade de linhas é a mesma para os dois arquivos"""
        for i in range(len(self.document1)):
            self.lista1.append(self.document1[i].text)
        for i in range(len(self.document2)):
            self.lista2.append(self.document2[i].text)
        
        if len(self.lista1) == len(self.lista2):
            return(self.compara_arquivos())
        else:
            return(self.verifica_linhas())

    def compara_arquivos(self):
        """Compara os dois arquivos, linha a linha, dada que a quantidade de linhas é a mesma para os dois"""
        arq1 = []
        arq2 = []
        for i,j in zip(self.lista1, self.lista2):
            if i != j:
                arq1.append(i)
                arq2.append(j)
        if len(arq1) == 0 and len(arq2) == 0:
            return('Os arquivos são iguais!')
        else:
            data = {'Arquivo1':arq1, 'Arquivo2': arq2}    
            df = pd.DataFrame(data, columns=['Arquivo1', 'Arquivo2'])
            return(df)
    
    def verifica_linhas(self):
        """Retorna as linhas que diferem entre si, e as linhas únicas de um dos documentos."""
        if len(self.lista1) > len(self.lista2):
            """Completar a diferença entre as listas com valores vazios"""
            for i in range(len(self.lista2), len(self.lista1)):
                self.lista2.append('Linha vazia')
                """Retornar a função compara_arquivos com o mesmo tamanho para ambas as listas"""
            return(self.compara_arquivos())
        
        if len(self.lista2) > len(self.lista1):
            """Completar a diferença entre as listas com valores vazios"""
            for i in range(len(self.lista1), len(self.lista2)):
                self.lista1.append('Linha vazia')
                """Retornar a função compara_arquivos com o mesmo tamanho para ambas as listas"""
            return(self.compara_arquivos())